import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def markdown_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        return

    doc = Document()
    
    # Title Style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=0)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=2)
        elif line.startswith('---'):
            doc.add_page_break()
        elif line.startswith('|'):
            # Simple table handling (basic)
            if '---' in line: continue
            cells = [c.strip() for c in line.split('|') if c.strip()]
            if not cells: continue
            # For simplicity in this script, we'll just add as a bold paragraph if it's the first time
            p = doc.add_paragraph()
            p.add_run(" | ".join(cells)).bold = True
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('1. '):
            doc.add_paragraph(line[3:], style='List Number')
        else:
            doc.add_paragraph(line)

    doc.save(docx_path)
    print(f"Word document saved to {docx_path}")

if __name__ == "__main__":
    md_file = r"C:\Users\LENOVO\.gemini\antigravity\brain\9427a8c8-31e3-4824-a8fc-0855c9f1cbc9\case_study_report.md"
    docx_file = "Voice_Stress_Detection_Case_Study.docx"
    markdown_to_docx(md_file, docx_file)
